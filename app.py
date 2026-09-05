import streamlit as st
import pandas as pd
import io
import os
import glob
import re
from modulos.parser_samm import limpiar_reporte_samm
from modulos.generador_excel import generar_pdf_cliente, generar_pdf_interno_dino

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="Dinomontacargas - ERP Mantenimiento", 
    page_icon="🚜", 
    layout="wide"
)

# --- MENÚ LATERAL (ENRUTADOR) ---
with st.sidebar:
    st.header("🛠️ Módulos de Gestión")
    menu_seleccionado = st.radio(
        "Navegación:",
        [
            "📅 1. Gestión de Cronogramas", 
            "🛢️ 2. Predictivo de Horómetros", 
            "🚜 3. Directorio de Flota"
        ]
    )
    
    st.markdown("---")
    st.header("📥 Ingreso de Datos (Módulo 1)")
    st.info("Sube el reporte de SAMM (Detalle_Visitas).")
    archivo_samm = st.file_uploader("Cargar SAMM", type=["xls", "xlsx"], key="samm")

# =====================================================================
# --- CEREBRO GLOBAL: LECTURA DE LA BASE MAESTRA ---
# =====================================================================
if 'df_base_maestra' not in st.session_state:
    if not os.path.exists("datos_samm"):
        st.error("⚠️ La carpeta 'datos_samm' no existe. Por favor, créala y mete tu base de equipos ahí.")
        st.stop()
        
    archivos = glob.glob("datos_samm/*.xls*")
    archivos_maestros = [f for f in archivos if "Control_Mantenimiento" not in f]
    
    if not archivos_maestros:
        st.error("⚠️ La carpeta 'datos_samm' está vacía. Coloca tu archivo de equipos ahí.")
        st.stop()
        
    try:
        ruta_maestro = archivos_maestros[0]
        df_maestro = pd.read_excel(ruta_maestro) 
        st.session_state['df_base_maestra'] = df_maestro
    except Exception as e:
        st.error(f"⚠️ Error crítico al leer la Base Maestra. Detalle: {e}")
        st.stop()

# =====================================================================
# 🟩 MÓDULO 1: GESTIÓN DE CRONOGRAMAS 
# =====================================================================
if menu_seleccionado == "📅 1. Gestión de Cronogramas":
    st.title("📅 Gestión de Cronogramas y Auditoría")
    st.markdown("---")
    
    if archivo_samm is not None:
        try:
            df_maestro_actual = st.session_state['df_base_maestra'].copy()
            
            # --- NORMALIZADOR UNIVERSAL DE IDs (FRANCOTIRADOR DE CORCHETES) ---
            def normalizar_id_universal(val):
                val_str = str(val).strip()
                match = re.search(r'\[\s*(\d+)\s*\]', val_str)
                if match: return match.group(1)
                if val_str.endswith('.0'): return val_str[:-2]
                return val_str
                
            df_maestro_actual['equipo_clean'] = df_maestro_actual['equipo'].apply(normalizar_id_universal)
            
            # --- DICCIONARIOS BLINDADOS ---
            claves_maestro = df_maestro_actual['equipo_clean'].astype(str).str.strip()
            dict_tercero = dict(zip(claves_maestro, df_maestro_actual['Tercero']))
            
            # PRIORIDAD ESTRICTA A LA SUCURSAL
            col_ubi = 'Sucursal' if 'Sucursal' in df_maestro_actual.columns else 'sucursal' if 'sucursal' in df_maestro_actual.columns else 'Ubicacion'
            dict_sucursal = dict(zip(claves_maestro, df_maestro_actual[col_ubi].astype(str).str.strip()))

            if 'df_master' not in st.session_state or st.session_state.get('nombre_archivo') != archivo_samm.name:
                df_crudo = limpiar_reporte_samm(archivo_samm)
                
                # --- LIMPIEZA GLOBAL INMEDIATA DE SERIALES ---
                df_crudo['Equipo'] = df_crudo['Equipo'].apply(normalizar_id_universal)
                
                # INICIALIZAR MEMORIA DE IGNORADOS
                if 'equipos_ignorados' not in st.session_state:
                    st.session_state['equipos_ignorados'] = set()

                # --- AUDITORÍA ESTRICTA (CLIENTE SAMM vs TERCERO MAESTRO) ---
                def auditar_contra_maestro(row):
                    equipo = str(row['Equipo']).strip() 
                    if equipo in st.session_state['equipos_ignorados']:
                        return "VERDE"
                        
                    cliente_samm = str(row['Cliente']).upper().strip()
                    tercero_maestro_raw = dict_tercero.get(equipo)
                    
                    if pd.isna(tercero_maestro_raw) or str(tercero_maestro_raw).strip() == "" or str(tercero_maestro_raw).strip().upper() in ["NAN", "SIN_ASIGNAR"]: 
                        return "VERDE" 
                        
                    tercero_maestro = str(tercero_maestro_raw).upper().strip()
                    
                    if cliente_samm not in tercero_maestro and tercero_maestro not in cliente_samm: 
                        return "ROJO" 
                    return "VERDE"
                    
                df_crudo['Alerta_Auditoria'] = df_crudo.apply(auditar_contra_maestro, axis=1)

                # --- MOTOR DE AGRUPACIÓN LOGÍSTICA (SEMANA ISO + BALDES DE 4) ---
                def optimizar_fechas_por_sucursal_y_cupos(df):
                    df_opt = df.copy()
                    df_opt['Fecha_DT'] = pd.to_datetime(df_opt['Fecha_Visita'].astype(str).str.split(" ").str[0], dayfirst=True, errors='coerce')
                    
                    # Agrupar por la sucursal exacta del Excel Maestro
                    df_opt['Sucursal_Maestra'] = df_opt['Equipo'].apply(lambda x: dict_sucursal.get(x, "SIN_SUCURSAL")).astype(str).str.upper().str.strip()
                    df_opt['Semana'] = df_opt['Fecha_DT'].dt.isocalendar().week
                    df_opt['Año'] = df_opt['Fecha_DT'].dt.isocalendar().year
                    
                    cambios_realizados = 0
                    grupos = df_opt.groupby(['Cliente', 'Sucursal_Maestra', 'Año', 'Semana'])
                    
                    for (cliente, sucursal, anio, semana), grupo in grupos:
                        if pd.isna(semana) or sucursal in ["SIN_SUCURSAL", "NAN", ""]: continue
                        
                        if len(grupo) > 1:
                            dias_disponibles = sorted(grupo['Fecha_DT'].dropna().unique())
                            if not dias_disponibles: continue
                            
                            grupo_ordenado = grupo.sort_values('Fecha_DT')
                            idx_dia = 0
                            cupo_actual = 0
                            
                            for idx, row in grupo_ordenado.iterrows():
                                if pd.isna(row['Fecha_DT']): continue
                                
                                dia_asignar = dias_disponibles[idx_dia]
                                
                                if row['Fecha_DT'] != dia_asignar:
                                    df_opt.at[idx, 'Fecha_Visita'] = dia_asignar.strftime("%d/%m/%Y")
                                    df_opt.at[idx, 'Fecha_DT'] = dia_asignar
                                    cambios_realizados += 1
                                    
                                cupo_actual += 1
                                # Llena el balde de 4, pasa al siguiente día que YA ESTABA AGENDADO
                                if cupo_actual >= 4:
                                    cupo_actual = 0
                                    if idx_dia < len(dias_disponibles) - 1:
                                        idx_dia += 1

                    df_opt = df_opt.drop(columns=['Fecha_DT', 'Sucursal_Maestra', 'Semana', 'Año'])
                    return df_opt, cambios_realizados

                df_crudo, total_optimizados = optimizar_fechas_por_sucursal_y_cupos(df_crudo)
                
                if total_optimizados > 0:
                    st.toast(f"🚜 Logística Activa: {total_optimizados} visitas organizadas en baldes de 4 por semana.", icon="🚜")

                st.session_state['df_master'] = df_crudo
                st.session_state['nombre_archivo'] = archivo_samm.name

            df_limpio = st.session_state['df_master']
            
            total_previstos = len(df_limpio)
            if total_previstos > 0:
                sin_ot = len(df_limpio[df_limpio['Color_Semantico'] == 'Rojo'])
                en_proceso = len(df_limpio[df_limpio['Color_Semantico'] == 'Amarillo'])
                finalizadas = len(df_limpio[df_limpio['Color_Semantico'] == 'Verde'])
            else:
                sin_ot, en_proceso, finalizadas = 0, 0, 0

            st.subheader("📊 Resumen General de Operaciones")
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Total Visitas", f"{total_previstos}")
            col2.metric("🔴 Sin OT (Rojo)", f"{sin_ot}")
            col3.metric("🟡 En Proceso (Amarillo)", f"{en_proceso}")
            col4.metric("🟢 Finalizadas (Verde)", f"{finalizadas}")
            st.markdown("---")
            
            tab_buscador, tab_datos, tab_cronograma = st.tabs([
                "🔍 Buscador de Flota", "📋 Base de Datos General", "📅 Generar Cronogramas y Auditoría"
            ])
            
            with tab_buscador:
                st.subheader("Radiografía por Equipo")
                lista_equipos = sorted(df_limpio['Equipo'].unique())
                equipo_buscado = st.selectbox("Selecciona el Equipo:", options=["Seleccionar..."] + lista_equipos)
                if equipo_buscado != "Seleccionar...":
                    df_equipo = df_limpio[df_limpio['Equipo'] == equipo_buscado].sort_values(by="Fecha_Visita")
                    c1, c2, c3 = st.columns(3)
                    c1.info(f"**🏢 Cliente:**\n{df_equipo['Cliente'].iloc[0]}")
                    c2.info(f"**📍 Ubicación:**\n{df_equipo['Sucursal'].iloc[0]}")
                    c3.info(f"**🔧 Visitas:**\n{len(df_equipo)}")
                    st.dataframe(df_equipo[['Fecha_Visita', 'Mantenimiento', 'OT', 'Estado']], use_container_width=True, hide_index=True)
            
            with tab_datos:
                st.subheader("Auditoría Global de la Flota (Filtro Comercial)")
                df_anomalias = df_limpio[df_limpio['Alerta_Auditoria'] == 'ROJO'].drop_duplicates(subset=['Equipo', 'Cliente']).copy()
                
                if not df_anomalias.empty:
                    st.warning(f"🚨 Hay {len(df_anomalias)} anomalías en toda la flota (Cliente en SAMM vs Tercero en Maestro).")
                    st.dataframe(df_anomalias[['Equipo', 'Cliente', 'Sucursal', 'Estado', 'Mantenimiento']], hide_index=True, use_container_width=True)
                    
                    buffer = io.BytesIO()
                    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                        df_anomalias.to_excel(writer, sheet_name='Anomalias', index=False)
                    st.download_button("📥 Descargar Reporte Global", buffer.getvalue(), "Anomalias.xlsx", "application/vnd.ms-excel", type="primary")
                else:
                    st.success("✅ Toda la flota operativa coincide correctamente con el Excel Maestro.")
                    
                st.write("---")
                
                # --- NUEVA ALERTA GLOBAL: VISITAS DUPLICADAS ---
                st.subheader("📌 Equipos con Múltiples Visitas el Mismo Día")
                df_global_multiples = df_limpio.groupby(['Cliente', 'Equipo', 'Fecha_Visita']).size().reset_index(name='Cantidad_Visitas')
                df_global_multiples = df_global_multiples[df_global_multiples['Cantidad_Visitas'] > 1]
                
                if not df_global_multiples.empty:
                    st.info(f"Se detectaron {len(df_global_multiples)} casos de equipos con más de una visita programada para la misma fecha.")
                    st.dataframe(df_global_multiples, hide_index=True)
                else:
                    st.success("✅ No hay visitas duplicadas para el mismo día en la flota.")
                    
                st.write("---")
    
                st.subheader("🚨 Equipos Sin Cronograma de Mantenimiento")
                
                equipos_en_samm = df_limpio['Equipo'].unique()
                df_faltantes = df_maestro_actual[~df_maestro_actual['equipo_clean'].isin(equipos_en_samm)]
                df_faltantes_mostrar = df_faltantes[['equipo_clean', 'Tercero', col_ubi, 'Modelo', 'Horometro Actual']].dropna(subset=['Tercero'])
                df_faltantes_mostrar.rename(columns={'equipo_clean': 'equipo'}, inplace=True)
                
                if not df_faltantes_mostrar.empty:
                    st.warning(f"{len(df_faltantes_mostrar)} equipos sin visita programada.")
                    st.dataframe(df_faltantes_mostrar, hide_index=True)
                    df_word = df_faltantes_mostrar.drop(columns=['Horometro Actual'], errors='ignore')
                    
                    def generar_word_faltantes(df):
                        from docx import Document 
                        doc = Document()
                        doc.add_heading('Equipos Sin Proyección', level=1)
                        doc.add_paragraph(f'Se reportan {len(df)} equipos activos sin visita.')
                        tabla = doc.add_table(rows=1, cols=len(df.columns))
                        tabla.style = 'Table Grid'
                        hdr_cells = tabla.rows[0].cells
                        for i, col in enumerate(df.columns): hdr_cells[i].text = str(col).upper()
                        for _, fila in df.iterrows():
                            row_cells = tabla.add_row().cells
                            for i, val in enumerate(fila): row_cells[i].text = str(val)
                        buf = io.BytesIO()
                        doc.save(buf)
                        return buf.getvalue()

                    st.download_button("📄 Descargar Informe (Word)", generar_word_faltantes(df_word), "Equipos_Faltantes.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
                    
                    st.write("---")
                    with st.expander("🛠️ Asignar a un Cronograma (Modo Manual Temporal)"):
                        st.info("Inyecta equipos huérfanos directamente al cronograma de un cliente en memoria.")
                        
                        col_eq, col_cli = st.columns(2)
                        lista_huerfanos = sorted(df_faltantes_mostrar['equipo'].unique())
                        
                        eq_manual = col_eq.selectbox("1. Equipo sin proyección:", lista_huerfanos)
                        
                        clientes_existentes = sorted(df_limpio['Cliente'].unique())
                        tipo_cli = col_cli.radio("2. Asignar a:", ["Cliente Existente", "Crear Nuevo Cliente"])
                        if tipo_cli == "Cliente Existente":
                            cli_destino = col_cli.selectbox("Selecciona el cliente:", clientes_existentes)
                        else:
                            cli_destino = col_cli.text_input("Escribe el nuevo cliente:")
                            
                        st.markdown("**3. Fechas de Visita (Formato: DD/MM/YYYY)**")
                        num_visitas = st.number_input("¿Cuántas visitas deseas programar?", min_value=1, max_value=6, value=1)
                        
                        cols_fechas = st.columns(num_visitas)
                        fechas_ingresadas = []
                        
                        for i in range(num_visitas):
                            label = f"Visita {i+1} {'*' if i == 0 else ''}"
                            f = cols_fechas[i].text_input(label, key=f"fecha_manual_{i}", placeholder="Ej: 15/08/2026")
                            fechas_ingresadas.append(f)
                            
                        if st.button("💾 Guardar y Asignar Equipo", type="primary"):
                            if not cli_destino or cli_destino.strip() == "":
                                st.error("⚠️ El nombre del cliente es obligatorio.")
                            elif not fechas_ingresadas[0] or fechas_ingresadas[0].strip() == "":
                                st.error("⚠️ Debes ingresar obligatoriamente la primera fecha.")
                            else:
                                info_eq = df_faltantes_mostrar[df_faltantes_mostrar['equipo'] == eq_manual].iloc[0]
                                sucursal_eq = info_eq.iloc[2] if pd.notna(info_eq.iloc[2]) else "SIN SUCURSAL"
                                
                                nuevas_filas = []
                                for fecha in fechas_ingresadas:
                                    if fecha and fecha.strip() != "":
                                        nueva_fila = {
                                            'Equipo': eq_manual,
                                            'Cliente': cli_destino.upper().strip(),
                                            'Sucursal': sucursal_eq,
                                            'Fecha_Visita': fecha.strip(), 
                                            'Estado': 'PROGRAMADO MANUAL',
                                            'Mantenimiento': 'PREVENTIVO',
                                            'Color_Semantico': 'Amarillo',
                                            'Alerta_Auditoria': 'VERDE'
                                        }
                                        nuevas_filas.append(nueva_fila)
                                
                                df_nuevas = pd.DataFrame(nuevas_filas)
                                st.session_state['df_master'] = pd.concat([st.session_state['df_master'], df_nuevas], ignore_index=True)
                                st.success(f"✅ ¡Equipo {eq_manual} inyectado exitosamente al contrato {cli_destino.upper()}!")
                                st.rerun()
                else:
                    st.success("✅ Todos tienen cronograma.")

            with tab_cronograma:
                st.subheader("Auditoría y Automatización de PDF (Clientes)")
                clientes_unicos = sorted(df_limpio['Cliente'].unique())
                opciones_selector = [f"👀 {c}" if 'ROJO' in df_limpio[df_limpio['Cliente'] == c]['Alerta_Auditoria'].values else c for c in clientes_unicos]
                
                # --- MEMORIA DE POSICIÓN ---
                indice_guardado = 0
                if 'ultimo_cliente' in st.session_state:
                    for i, op in enumerate(opciones_selector):
                        if op.replace("👀 ", "") == st.session_state['ultimo_cliente']:
                            indice_guardado = i
                            break
                            
                cliente_raw = st.selectbox("Cliente:", opciones_selector, index=indice_guardado)
                cliente_seleccionado = cliente_raw.replace("👀 ", "")
                st.session_state['ultimo_cliente'] = cliente_seleccionado
                # ---------------------------
                
                df_filtrado = df_limpio[df_limpio['Cliente'] == cliente_seleccionado]
                alertas_rojas = df_filtrado[df_filtrado['Alerta_Auditoria'] == 'ROJO']
                
                if not alertas_rojas.empty:
                    st.warning("👀 REVISIÓN SUGERIDA: Conflicto entre Contrato SAMM y Tercero Maestro.")
                    
                    df_alertas_cli = alertas_rojas[['Equipo', 'Cliente', 'Sucursal']].drop_duplicates().copy()
                    df_alertas_cli.insert(0, '✅ Seleccionar', False)
                    
                    st.info("💡 Selecciona los equipos en la tabla y elige la acción abajo.")
                    df_editado_cli = st.data_editor(df_alertas_cli, hide_index=True, use_container_width=True, disabled=['Equipo', 'Cliente', 'Sucursal'])
                    df_seleccionados = df_editado_cli[df_editado_cli['✅ Seleccionar'] == True]
                    
                    # --- PANEL DE ACCIÓN RÁPIDA (DESPLEGABLES ESTABLES) ---
                    st.markdown("### 🛠️ Panel de Acción Rápida")
                    if not df_seleccionados.empty:
                        equipos_afectados = df_seleccionados['Equipo'].tolist()
                        
                        def ejecutar_reasignacion(nuevo_cliente_destino):
                            mascara_mover = st.session_state['df_master']['Equipo'].isin(equipos_afectados)
                            st.session_state['df_master'].loc[mascara_mover, 'Cliente'] = nuevo_cliente_destino
                            
                            def re_auditar(row):
                                eq_s = str(row['Equipo']).strip()
                                if eq_s in st.session_state.get('equipos_ignorados', set()): return "VERDE"
                                cliente_samm = str(row['Cliente']).upper().strip()
                                tercero_maestro_raw = dict_tercero.get(eq_s)
                                if pd.isna(tercero_maestro_raw) or str(tercero_maestro_raw).strip() == "" or str(tercero_maestro_raw).strip().upper() in ["NAN", "SIN_ASIGNAR"]: 
                                    return "VERDE"
                                tercero_maestro = str(tercero_maestro_raw).upper().strip()
                                if cliente_samm not in tercero_maestro and tercero_maestro not in cliente_samm: return "ROJO"
                                return "VERDE"
                                
                            st.session_state['df_master']['Alerta_Auditoria'] = st.session_state['df_master'].apply(re_auditar, axis=1)
                            st.rerun()

                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            if st.button("✅ Ignorar y Aprobar (Lista Blanca)", type="primary", use_container_width=True):
                                if 'equipos_ignorados' not in st.session_state:
                                    st.session_state['equipos_ignorados'] = set()
                                st.session_state['equipos_ignorados'].update(equipos_afectados) 
                                mascara = st.session_state['df_master']['Equipo'].isin(equipos_afectados)
                                st.session_state['df_master'].loc[mascara, 'Alerta_Auditoria'] = 'VERDE'
                                st.rerun()
                                
                        with col2:
                            with st.expander("🔄 Reasignar a Cliente Existente"):
                                dst_existente = st.selectbox("Selecciona destino:", clientes_unicos, label_visibility="collapsed", key="sel_dst")
                                if st.button("Confirmar Reasignación", use_container_width=True):
                                    if dst_existente and dst_existente.strip() != "":
                                        ejecutar_reasignacion(dst_existente.strip().upper())
                                    else:
                                        st.error("⚠️ Selecciona un destino.")
                                        
                        with col3:
                            with st.expander("📄 Crear Nuevo Contrato"):
                                dst_nuevo = st.text_input("Nombre del contrato:", label_visibility="collapsed", placeholder="Ej: TRACTOCAR CARTAGENA", key="txt_dst")
                                if st.button("Confirmar Creación", use_container_width=True):
                                    if dst_nuevo and dst_nuevo.strip() != "":
                                        ejecutar_reasignacion(dst_nuevo.strip().upper())
                                    else:
                                        st.error("⚠️ Escribe un nombre válido.")
                    else:
                        st.info("👈 Selecciona al menos un equipo en la tabla superior para activar las opciones.")
                            
                st.write("---")
                
                # --- NUEVA ALERTA LOCAL: VISITAS DUPLICADAS (POR CLIENTE) ---
                df_visitas_multiples = df_filtrado.groupby(['Equipo', 'Fecha_Visita']).size().reset_index(name='Cantidad_Visitas')
                df_visitas_multiples = df_visitas_multiples[df_visitas_multiples['Cantidad_Visitas'] > 1]
                
                if not df_visitas_multiples.empty:
                    st.info("📌 **ATENCIÓN:** Los siguientes equipos tienen múltiples visitas el mismo día. (En el PDF se purgarán y mostrará 1 sola fecha).")
                    st.dataframe(df_visitas_multiples, hide_index=True)
                    
                # =========================================================
                # NUEVO: VISTA PREVIA Y SEPARADOR DE CRONOGRAMAS
                # =========================================================
                st.write("---")
                st.subheader(f"📋 Vista Previa del Cronograma: {cliente_seleccionado}")
                st.info("Revisa los equipos que saldrán en este PDF. Si necesitas separar este cronograma en dos contratos distintos, marca las casillas y muévelos.")
                
                df_vista_previa = df_filtrado[['Equipo', 'Sucursal', 'Fecha_Visita', 'Mantenimiento']].copy()
                df_vista_previa.insert(0, '✅ Separar', False)
                
                df_editado_prev = st.data_editor(df_vista_previa, hide_index=True, use_container_width=True, disabled=['Equipo', 'Sucursal', 'Fecha_Visita', 'Mantenimiento'])
                equipos_a_separar = df_editado_prev[df_editado_prev['✅ Separar'] == True]['Equipo'].tolist()
                
                if equipos_a_separar:
                    st.markdown("### ✂️ Separar Equipos Seleccionados")
                    
                    def ejecutar_separacion(nuevo_cliente_destino):
                        mascara_sep = st.session_state['df_master']['Equipo'].isin(equipos_a_separar)
                        st.session_state['df_master'].loc[mascara_sep, 'Cliente'] = nuevo_cliente_destino
                        
                        def re_auditar_sep(row):
                            eq_s = str(row['Equipo']).strip()
                            if eq_s in st.session_state.get('equipos_ignorados', set()): return "VERDE"
                            cliente_samm = str(row['Cliente']).upper().strip()
                            tercero_maestro_raw = dict_tercero.get(eq_s)
                            if pd.isna(tercero_maestro_raw) or str(tercero_maestro_raw).strip() == "" or str(tercero_maestro_raw).strip().upper() in ["NAN", "SIN_ASIGNAR"]: return "VERDE"
                            tercero_maestro = str(tercero_maestro_raw).upper().strip()
                            if cliente_samm not in tercero_maestro and tercero_maestro not in cliente_samm: return "ROJO"
                            return "VERDE"
                            
                        st.session_state['df_master']['Alerta_Auditoria'] = st.session_state['df_master'].apply(re_auditar_sep, axis=1)
                        st.rerun()

                    c1_sep, c2_sep = st.columns(2)
                    with c1_sep:
                        with st.expander("🔄 Mover a Cliente Existente"):
                            dst_ex = st.selectbox("Destino:", clientes_unicos, label_visibility="collapsed", key="sel_dst_ex")
                            if st.button("Confirmar Traslado", use_container_width=True, key="btn_ex"):
                                if dst_ex and dst_ex.strip() != "": ejecutar_separacion(dst_ex.strip().upper())
                                else: st.error("⚠️ Selecciona un destino.")
                    with c2_sep:
                        with st.expander("📄 Mover a Nuevo Contrato"):
                            dst_nu = st.text_input("Nombre del contrato:", label_visibility="collapsed", placeholder="Ej: CLIENTE - SEDE NORTE", key="sel_dst_nu")
                            if st.button("Confirmar Creación", use_container_width=True, key="btn_nu"):
                                if dst_nu and dst_nu.strip() != "": ejecutar_separacion(dst_nu.strip().upper())
                                else: st.error("⚠️ Escribe un nombre válido.")
                                
                st.write("---")
                
                st.subheader("📝 Novedades y Observaciones del Mes")
                texto_novedad = st.text_area(
                    "Agrega una nota personalizada para este cliente (opcional):", 
                    placeholder="Ej: Durante este mes se observó un desgaste irregular en las llantas del equipo 408..."
                )
                
                lista_combustion = []
                archivo_tracker = "datos_samm/Control_Mantenimiento.xlsx"
                if os.path.exists(archivo_tracker):
                    try:
                        df_tr = pd.read_excel(archivo_tracker)
                        col = 'EQUIPO' if 'EQUIPO' in df_tr.columns else 'Equipo'
                        lista_combustion = df_tr[col].astype(str).str.strip().tolist()
                    except: pass

                pdf_bytes = generar_pdf_cliente(df_filtrado, cliente_seleccionado, texto_novedad, lista_combustion)
                
                st.download_button(f"⬇️ Descargar PDF - {cliente_seleccionado}", pdf_bytes, f"Cronograma_{cliente_seleccionado}.pdf", "application/pdf", type="primary")

                st.markdown("---")
                st.subheader("🚜 Planificador Logístico Interno (Mantenimiento Dino)")
                st.info("Calcula las horas técnicas del mes, evaluando restricciones FÍSICAS (Sucursal).")
                
                palabras_excluidas = ['cartagena', 'ajover', 'bogota', 'altipal', 'cerete', 'serete', 'soberana']
                patron_exclusion = '|'.join(palabras_excluidas)
                
                mascara_sucursal = df_limpio['Sucursal'].astype(str).str.contains(patron_exclusion, case=False, na=False)
                df_ruta_dino = df_limpio[~mascara_sucursal].copy()
                
                if not df_ruta_dino.empty:
                    total_visitas_ruta = len(df_ruta_dino)
                    horas_estimadas = total_visitas_ruta * 2
                    
                    c1, c2 = st.columns(2)
                    c1.metric("📍 Visitas en Ruta (Excluyendo foráneos)", total_visitas_ruta)
                    c2.metric("⏱️ Horas Técnicas Estimadas (2h x Eq)", f"{horas_estimadas} hrs")
                    
                    pdf_interno_bytes = generar_pdf_interno_dino(df_ruta_dino, horas_estimadas)
                    
                    st.download_button(
                        label="⚙️ Descargar Cronograma Interno (PDF)",
                        data=pdf_interno_bytes,
                        file_name="Cronograma_Interno_Dino.pdf",
                        mime="application/pdf"
                    )
                    
                    with st.expander("Ver lista de equipos incluidos en esta ruta"):
                        st.dataframe(df_ruta_dino[['Cliente', 'Equipo', 'Sucursal', 'Fecha_Visita']].sort_values(by=['Sucursal', 'Fecha_Visita']), hide_index=True)
                else:
                    st.warning("No se encontraron equipos para esta ruta después de aplicar las exclusiones.")

        except Exception as e:
            st.error(f"Error crítico al procesar el archivo. Detalles: {e}")
    else:
        st.info("👈 Sube el reporte de SAMM para gestionar los cronogramas.")

# =====================================================================
# 🟦 MÓDULO 2: PREDICTIVO DE HORÓMETROS (Aceites y Filtros)
# =====================================================================
elif menu_seleccionado == "🛢️ 2. Predictivo de Horómetros":
    st.title("🛢️ Control Predictivo de Mantenimientos")
    st.markdown("---")
    
    df_maestro = st.session_state['df_base_maestra'].copy()
    
    def limpiar_horometro_base(val):
        if pd.isna(val): return None
        val_str = str(val).strip().replace(',', '.')
        try:
            num = float(val_str)
            while num > 40000:
                num = num / 10.0
            return num
        except:
            return None
            
    df_maestro['Horometro Actual'] = df_maestro['Horometro Actual'].apply(limpiar_horometro_base)
    df_maestro = df_maestro.dropna(subset=['Horometro Actual'])
    
    def normalizar_id_universal(val):
        val_str = str(val).strip()
        match = re.search(r'\[\s*(\d+)\s*\]', val_str)
        if match: return match.group(1)
        if val_str.endswith('.0'): return val_str[:-2]
        return val_str

    df_maestro['equipo_str'] = df_maestro['equipo'].apply(normalizar_id_universal)
    
    archivo_tracker = "datos_samm/Control_Mantenimiento.xlsx"
    
    if os.path.exists(archivo_tracker):
        df_tracker_crudo = pd.read_excel(archivo_tracker)
        df_tracker = df_tracker_crudo.copy()
        
        if 'EQUIPO' in df_tracker.columns:
            df_tracker.rename(columns={'EQUIPO': 'Equipo'}, inplace=True)
        if 'Horometro de cambio deaceite' in df_tracker.columns:
            df_tracker.rename(columns={'Horometro de cambio deaceite': 'Ultimo_Mantenimiento'}, inplace=True)
            
        if 'ESTADO INSUMOS' not in df_tracker.columns:
            df_tracker['Estado_Insumos'] = "Al día"
        else:
            df_tracker.rename(columns={'ESTADO INSUMOS': 'Estado_Insumos'}, inplace=True)
            df_tracker['Estado_Insumos'] = df_tracker['Estado_Insumos'].fillna("Al día")
            
        if 'Horometro_Base_Ciclo' not in df_tracker.columns:
            df_tracker['Horometro_Base_Ciclo'] = df_tracker['Ultimo_Mantenimiento']
            
        df_tracker['Equipo'] = df_tracker['Equipo'].apply(normalizar_id_universal)
        df_tracker = df_tracker[~df_tracker['Equipo'].str.upper().isin(['NAN', 'NONE', 'NA', ''])]
        df_maestro = df_maestro[~df_maestro['equipo_str'].str.upper().isin(['NAN', 'NONE', 'NA', ''])]
    else:
        st.warning("⚠️ No se encontró el archivo 'Control_Mantenimiento.xlsx' en 'datos_samm'.")
        st.stop()
        
    # PRIORIDAD ESTRICTA A LA SUCURSAL
    col_ubi_pred = 'Sucursal' if 'Sucursal' in df_maestro.columns else 'sucursal' if 'sucursal' in df_maestro.columns else 'Ubicacion'
    
    df_predictivo = pd.merge(
        df_tracker, 
        df_maestro[['equipo_str', 'Tercero', col_ubi_pred, 'Modelo', 'Horometro Actual']], 
        left_on='Equipo', 
        right_on='equipo_str', 
        how='inner'
    )
    
    df_predictivo['Ultimo_Mantenimiento'] = pd.to_numeric(df_predictivo['Ultimo_Mantenimiento'], errors='coerce').fillna(0)
    df_predictivo['Horometro_Base_Ciclo'] = pd.to_numeric(df_predictivo['Horometro_Base_Ciclo'], errors='coerce').fillna(df_predictivo['Ultimo_Mantenimiento'])
    
    def corregir_coma_flotante(row):
        actual_original = row['Horometro Actual']
        ultimo = row['Ultimo_Mantenimiento']
        if pd.isna(actual_original) or pd.isna(ultimo) or ultimo == 0: return actual_original
            
        actual = actual_original
        if (actual - ultimo) > 1000:
            for _ in range(2): 
                prueba = actual / 10.0
                if prueba >= (ultimo - 500) and (prueba - ultimo) < 1500: return prueba 
                actual = prueba
        return actual_original
        
    df_predictivo['Horometro Actual'] = df_predictivo.apply(corregir_coma_flotante, axis=1)

    df_predictivo['Proximo_Mantenimiento'] = df_predictivo['Ultimo_Mantenimiento'] + 250
    df_predictivo['Horas_Faltantes'] = df_predictivo['Proximo_Mantenimiento'] - df_predictivo['Horometro Actual']
    
    def calcular_nivel(row):
        faltan = row['Horas_Faltantes']
        if faltan < -50: return "NIVEL MÁXIMO (Requiere Cambio TOTAL por Atraso)"
        horas_acumuladas = row['Proximo_Mantenimiento'] - row['Horometro_Base_Ciclo']
        if horas_acumuladas <= 0: return "NIVEL 1 (Filtros Básicos Motor)"
        ciclo_exacto = round(horas_acumuladas / 250) * 250
        if ciclo_exacto % 2000 == 0: return "NIVEL 4 (Básico + Diferencial + Hidráulico)"
        if ciclo_exacto % 1000 == 0: return "NIVEL 3 (Básico + Caja Int + Correa)"
        if ciclo_exacto % 500 == 0: return "NIVEL 2 (Básico + Caja Ext + Frenos)"
        return "NIVEL 1 (Filtros Básicos Motor)"
        
    df_predictivo['Tipo_Mantenimiento'] = df_predictivo.apply(calcular_nivel, axis=1)
    
    def semaforo(faltan):
        if faltan <= 0: return "🔴 VENCIDO (Urgente)"
        if faltan <= 30: return "🟡 PREVENTIVO (Pedir Insumos)"
        return "🟢 ÓPTIMO"
        
    df_predictivo['Alerta'] = df_predictivo['Horas_Faltantes'].apply(semaforo)
    
    st.subheader("🚨 Panel de Alertas y Trámites")
    df_alertas = df_predictivo[df_predictivo['Alerta'].str.contains('🔴|🟡')].sort_values(by='Horas_Faltantes')
    
    if not df_alertas.empty:
        columnas_ver = ['Equipo', 'Tercero', 'Horometro Actual', 'Horas_Faltantes', 'Tipo_Mantenimiento', 'Alerta', 'Estado_Insumos']
        st.dataframe(df_alertas[columnas_ver], hide_index=True)
        
        buffer_alertas = io.BytesIO()
        with pd.ExcelWriter(buffer_alertas, engine='xlsxwriter') as writer:
            df_alertas[columnas_ver].to_excel(writer, sheet_name='Solicitud_Filtros', index=False)
        st.download_button(
            label="📥 Descargar Reporte para Compras (Excel)",
            data=buffer_alertas.getvalue(),
            file_name="Solicitud_Insumos_Mantenimiento.xlsx",
            mime="application/vnd.ms-excel"
        )
        
        st.markdown("---")
        st.subheader("⚙️ Gestión de Estado del Mantenimiento")
        c1, c2, c3 = st.columns([2, 2, 1])
        with c1: equipo_revisado = st.selectbox("1. Selecciona el equipo:", df_alertas['Equipo'].tolist())
        with c2: nuevo_estado = st.selectbox("2. Actualizar estado a:", ["Solicitado", "Pendiente por instalar", "✅ Cambio BÁSICO Realizado", "🚨 Cambio TOTAL Realizado (Reseteo)"])
        with c3:
            st.write("")
            st.write("")
            if st.button("Guardar Estado", type="primary", use_container_width=True):
                if "Realizado" in nuevo_estado:
                    nuevo_horometro = df_predictivo.loc[df_predictivo['Equipo'] == equipo_revisado, 'Horometro Actual'].values[0]
                    df_tracker.loc[df_tracker['Equipo'] == equipo_revisado, 'Ultimo_Mantenimiento'] = nuevo_horometro
                    df_tracker.loc[df_tracker['Equipo'] == equipo_revisado, 'Estado_Insumos'] = "Al día"
                    if "TOTAL" in nuevo_estado:
                        df_tracker.loc[df_tracker['Equipo'] == equipo_revisado, 'Horometro_Base_Ciclo'] = nuevo_horometro
                        st.success(f"¡Reseteo exitoso! El equipo {equipo_revisado} inicia un nuevo ciclo.")
                    else:
                        st.success(f"Mantenimiento básico registrado.")
                    from datetime import datetime
                    if 'Fecha de cambio aceite' in df_tracker.columns:
                        df_tracker.loc[df_tracker['Equipo'] == equipo_revisado, 'Fecha de cambio aceite'] = datetime.now().strftime("%Y-%m-%d")
                else:
                    df_tracker.loc[df_tracker['Equipo'] == equipo_revisado, 'Estado_Insumos'] = nuevo_estado
                    st.success(f"Estado actualizado a: {nuevo_estado}")
                
                df_salida = df_tracker.rename(columns={'Equipo': 'EQUIPO', 'Ultimo_Mantenimiento': 'Horometro de cambio deaceite', 'Estado_Insumos': 'ESTADO INSUMOS'})
                df_salida.to_excel(archivo_tracker, index=False)
                st.rerun()
    else:
        st.success("✅ Toda la flota está en estado óptimo.")
        
    with st.expander("📊 Ver Flota Completa (Combustión) y Estados"):
        st.dataframe(df_predictivo[['Equipo', 'Tercero', 'Horometro Actual', 'Ultimo_Mantenimiento', 'Horometro_Base_Ciclo', 'Horas_Faltantes', 'Tipo_Mantenimiento', 'Estado_Insumos']], hide_index=True)

# =====================================================================
# 🚜 MÓDULO 3: DIRECTORIO DE FLOTA
# =====================================================================
elif menu_seleccionado == "🚜 3. Directorio de Flota":
    st.title("🚜 Directorio Global de Flota")
    st.markdown("---")
    
    df_dir = st.session_state['df_base_maestra'].copy()
    busqueda = st.text_input("🔍 Buscar por Número de Equipo, Cliente, Modelo o Sucursal:")
    
    if busqueda:
        df_dir = df_dir[df_dir.astype(str).apply(lambda x: x.str.contains(busqueda, case=False)).any(axis=1)]
        
    # PRIORIDAD ESTRICTA A LA SUCURSAL
    col_ubi_dir = 'Sucursal' if 'Sucursal' in df_dir.columns else 'sucursal' if 'sucursal' in df_dir.columns else 'Ubicacion'
    columnas_dir = ['equipo', 'Tercero', col_ubi_dir, 'Modelo', 'Horometro Actual']
    if 'Link_Ficha' in df_dir.columns: columnas_dir.append('Link_Ficha')
        
    st.dataframe(df_dir[columnas_dir], use_container_width=True, hide_index=True)
    
    buffer_dir = io.BytesIO()
    with pd.ExcelWriter(buffer_dir, engine='xlsxwriter') as writer:
        df_dir[columnas_dir].to_excel(writer, sheet_name='Directorio', index=False)
        
    st.download_button(
        label="📥 Descargar Directorio de Pantalla (Excel)",
        data=buffer_dir.getvalue(),
        file_name="Directorio_Dinomontacargas.xlsx",
        mime="application/vnd.ms-excel"
    )
